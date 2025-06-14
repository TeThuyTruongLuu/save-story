import os
import json
import requests
from bs4 import BeautifulSoup
from seleniumwire import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import subprocess
import zipfile
from urllib.parse import urljoin, urlparse
from flask import Flask, request, jsonify, send_file, send_from_directory
import re
import io
import shutil
from threading import Lock
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
import unicodedata
from selenium.common.exceptions import TimeoutException

app = Flask(__name__)
CORS(app)

os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"

selector_configs = {
    "lofter_author": {
        "post_list": ["div.block.article", "div.m-post", "div.block.photo", "div.ztstext.article", "div.m-postdtl", "div.m-post.m-post-photo"],
        "post_url": ["a[href*='.lofter.com/post/']"],
        "title": ["div.header", "h2"],
        "content": ["div.m-post.m-post-photo", "div.content", "div.wrap", "div.text", "div.txtcont", "div.section", "p", "div.ct", "div.section", "div.m-detail",
        ],
        "images": ["div.content img", "div.wrap img", "div.pic img", "div.m-detail img", "img.img", "img.pic"]
    },
    "lofter_tag": {
        "post_list": ["div.m-mlist"],
        "post_url": ["div.isayt a.isayc[href*='.lofter.com/post/']"],
        "title": ["div.m-icnt h2.tit"],
        "content": ["div.m-icnt div.cnt div.txt[style*='display:none']"],
        "images": ["div.m-icnt div.cnt div.txt[style*='display:none'] img"]
    },
    "ao3": {
        # cũng riêng, code sẵn rồi
    }
}

lofter_cookies = None
cookie_lock = Lock()
driver_lock = Lock()
driver = None

def setup_selenium():
    global driver
    with driver_lock:
        if driver is None or not driver.session_id:
            options = Options()
            options.add_argument("--headless")
            options.add_argument("--disable-gpu")
            options.add_argument("--no-sandbox")
            options.add_argument("--disable-dev-shm-usage")
            options.add_argument("--disable-blink-features=AutomationControlled")
            options.add_experimental_option("excludeSwitches", ["enable-automation"])
            driver = webdriver.Chrome(options=options)
        return driver

def setup_new_driver():
    options = Options()
    options.add_argument("--headless")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    driver = webdriver.Chrome(options=options)
    return driver

def load_lofter_cookies(driver, cookie_file="lofter_cookies.json"):
    global lofter_cookies
    with cookie_lock:
        if lofter_cookies is None:
            try:
                with open(cookie_file, "r") as f:
                    lofter_cookies = json.load(f)
            except FileNotFoundError:
                print("Không tìm thấy file cookie Lofter.")
                return False
        
        current_time = time.time()
        valid_cookies = [cookie for cookie in lofter_cookies if cookie.get("expires", -1) == -1 or cookie["expires"] > current_time]
        if not valid_cookies:
            print("Cookie Lofter đã hết hạn. Vui lòng cập nhật cookie mới.")
            return False
        
        driver.get("https://www.lofter.com")
        time.sleep(2)
        for cookie in valid_cookies:
            driver.add_cookie(cookie)
        
        driver.get("https://www.lofter.com")
        print("Đã tải cookie Lofter hợp lệ.")
        return True

def save_lofter_cookies(driver, cookie_file="lofter_cookies.json"):
    with open(cookie_file, "w") as f:
        json.dump(driver.get_cookies(), f)
    print("Saved Lofter cookies.")

def parse_by_selector(post_elem, selector_set, index_in_list=0):
    result = {}
    
    for sel in selector_set.get("content", []):
        elem = post_elem.select_one(sel)
        if elem:
            for span in elem.select("span.picNum"):
                span.decompose()

            result["content"] = elem.decode_contents().strip()
            result["content_text"] = elem.get_text(separator=" ", strip=True)
            break

    if "content" not in result:
        result["content"] = "Không tìm thấy nội dung"
        result["content_text"] = "Không tìm thấy nội dung"
    
    import html


    result["images"] = []
    for sel in selector_set.get("images", []):
        imgs = post_elem.select(sel)
        if imgs:
            for img in imgs:
                try:
                    # Ưu tiên bigimgsrc từ <a> nếu có
                    parent_a = img.find_parent("a")
                    if parent_a and parent_a.has_attr("bigimgsrc"):
                        img_url = parent_a["bigimgsrc"]
                    else:
                        img_url = img["src"] if img["src"].startswith("http") else "https:" + img["src"]
                    clean_url = img_url.split('?')[0].split('!')[0]
                    clean_url = html.unescape(clean_url)

                    result["images"].append(clean_url)
                except Exception as e:
                    import logging
                    logging.info(f"Lỗi khi parse ảnh: {e}")
            break

    for sel in selector_set.get("title", []):
        elem = post_elem.select_one(sel)
        if elem and elem.text.strip():
            result["title"] = elem.text.strip()
            break
    if "title" not in result or not result["title"].strip():
        first_img = post_elem.select_one("img")
        if first_img and first_img.get("alt"):
            result["title"] = first_img["alt"].strip()
        elif result.get("content_text", "").strip():
            result["title"] = result["content_text"][:40].strip()
        else:
            result["title"] = f"Bài số {index_in_list}"
    
    return result

def download_images(images, post_id):
    downloaded_images = []
    os.makedirs("temp/images", exist_ok=True)
        
    for i, img_url in enumerate(images):
        try:
            # Clean lại 1 lần nữa để chắc ăn
            clean_url = img_url.split('?')[0].split('!')[0]

            # Nếu là ảnh Discord thì set header để cứu
            headers = {}
            if 'discordapp.com' in clean_url or 'media.discordapp.net' in clean_url:
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:115.0) Gecko/20100101 Firefox/115.0",
                    "Referer": "https://discord.com/"
                }

            response = requests.get(clean_url, stream=True, timeout=10, headers=headers)
            response.raise_for_status()
            img_path = f"temp/images/{post_id}_{i}.jpg"
            with open(img_path, "wb") as f:
                f.write(response.content)
            downloaded_images.append(img_path)
        except Exception as e:
            print(f"Không thể tải hình ảnh {img_url}: {e}")
            downloaded_images.append(None)

    
    return downloaded_images

def fetch_ao3_works(url, driver):
    driver.get(url)
    works = []
    while True:
        soup = BeautifulSoup(driver.page_source, "html.parser")
        for work in soup.select("li.work"):
            title = work.select_one("h4.heading a").text.strip()
            work_url = urljoin("https://archiveofourown.org", work.select_one("h4.heading a")["href"])
            works.append({"title": title, "url": work_url})
        try:
            next_button = driver.find_element(By.CSS_SELECTOR, "a[rel='next']")
            next_button.click()
            WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.CSS_SELECTOR, "li.work")))
        except:
            break
    return works

def fetch_lofter_posts(url, driver, is_tag=False, max_posts=50, wait_time=1, continue_fetch=False, start_date_ts=None, end_date_ts=None):
    driver.set_page_load_timeout(60)

    try:
        driver.get(url)
    except TimeoutException:
        print(f"Timeout khi tải {url}, skip trang này")
        return []

    if not load_lofter_cookies(driver):
        return []

    try:
        driver.get(url)
    except TimeoutException:
        print(f"Timeout khi tải lại {url} sau load cookie, skip trang này")
        return []

    time.sleep(2)

    selector_set = selector_configs.get("lofter_tag" if is_tag else "lofter_author", {})
    post_selector = ", ".join(selector_set.get("post_list", ["div.m-post"]))

    print(f"Using post_selector: {post_selector}")

    posts = []
    last_height = driver.execute_script("return document.body.scrollHeight")

    if continue_fetch:
        driver.get(url)
        time.sleep(2)

    # Cuộn vài lần cho chắc ăn
    for _ in range(5):
        try:
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            WebDriverWait(driver, wait_time).until(
                lambda d: d.execute_script("return document.body.scrollHeight") != last_height
            )
            new_height = driver.execute_script("return document.body.scrollHeight")
            if new_height == last_height or len(posts) >= max_posts:
                break
            last_height = new_height
        except TimeoutException:
            break
        except Exception as e:
            print(f"Error during scroll: {e}")
            break

    soup = BeautifulSoup(driver.page_source, "html.parser")

    parsed_url = urlparse(url)
    author_domain = parsed_url.netloc
    print(f"Author domain: {author_domain}")

    print(f"Found {len(soup.select(post_selector))} post elements")
    for post_elem in soup.select(post_selector):
        if len(posts) >= max_posts:
            print(f"Đã đủ {max_posts} bài → dừng parse bài trong trang hiện tại")
            break

        post_url = ""
        for a_elem in post_elem.select("a[href]"):
            href = a_elem["href"]
            full_url = urljoin(url, href)
            if full_url.startswith(f"https://{author_domain}/post/"):
                post_url = full_url
                print(f"Found post_url: {post_url}")
                break

        if not post_url:
            print("Không tìm thấy post_url hợp lệ trong post_elem, skip bài này")
            continue

        time_elem = post_elem.select_one("span.time, div.time, div.info .time")
        post_time_ts = None
        if time_elem and time_elem.text.strip():
            try:
                post_time_struct = time.strptime(time_elem.text.strip(), "%Y-%m-%d")
                post_time_ts = time.mktime(post_time_struct)
            except Exception as e:
                print(f"Không parse được time: {time_elem.text.strip()}")

        if start_date_ts and post_time_ts and post_time_ts < start_date_ts:
            print(f"Bỏ qua post {post_url} vì trước ngày lọc")
            continue
        if end_date_ts and post_time_ts and post_time_ts > end_date_ts:
            print(f"Bỏ qua post {post_url} vì sau ngày lọc")
            continue

        parsed = parse_by_selector(post_elem, selector_set, index_in_list=len(posts))

        if post_url and post_url not in [p["url"] for p in posts]:
            posts.append({
                "title": parsed["title"],
                "url": post_url,
                "content": parsed["content"],
                "images": parsed["images"],
                "preview": parsed["content_text"][:20] + "..." if parsed["content_text"] else "",
                "time": time_elem.text.strip() if time_elem and time_elem.text.strip() else ""
            })
            print(f"Thêm bài thành công: {parsed['title']} - {post_url}")

    # Xử lý next page nếu có
    next_link = soup.select_one("a.next, a.next-page")
    remaining_posts = max_posts - len(posts)
    if next_link and remaining_posts > 0:
        next_url = urljoin(url, next_link["href"])
        print(f"Chuyển sang trang tiếp theo: {next_url} (còn cần lấy {remaining_posts} bài)")
        next_page_posts = fetch_lofter_posts(next_url, driver, is_tag, remaining_posts, wait_time, start_date_ts=start_date_ts, end_date_ts=end_date_ts)
        posts.extend(next_page_posts)
    else:
        if remaining_posts <= 0:
            print(f"Đã đủ {max_posts} bài tổng → không fetch tiếp trang sau")
        else:
            print("Không có next page → dừng fetch")

    print(f"Total posts found: {len(posts)}")
    for post in posts:
        print(f"Title: {post['title']} - URL: {post['url']} - Time: {post.get('time', '')}")

    return posts


def fetch_lofter_tag_posts(url, driver, max_posts=50, wait_time=1, continue_fetch=False, start_date_ts=None, end_date_ts=None, cookies_loaded=False):
    driver.set_page_load_timeout(60)

    try:
        driver.get(url)
    except TimeoutException:
        print(f"Timeout khi tải {url}, skip trang này")
        return []

    # Load cookies chỉ lần đầu
    if not cookies_loaded:
        if not load_lofter_cookies(driver):
            return []
        try:
            driver.get(url)
        except TimeoutException:
            print(f"Timeout khi tải lại {url} sau load cookie, skip trang này")
            return []
        time.sleep(2)
        cookies_loaded = True
    else:
        time.sleep(2)

    selector_set = selector_configs["lofter_tag"]
    post_selector = ", ".join(selector_set.get("post_list", ["div.m-mlist"]))

    print(f"[TAG] Using post_selector: {post_selector}")

    posts = []
    last_height = driver.execute_script("return document.body.scrollHeight")
    stop_fetch = False

    if continue_fetch:
        driver.get(url)
        time.sleep(2)

    # Cuộn vài lần
    for _ in range(5):
        try:
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            WebDriverWait(driver, wait_time).until(
                lambda d: d.execute_script("return document.body.scrollHeight") != last_height
            )
            new_height = driver.execute_script("return document.body.scrollHeight")
            if new_height == last_height or len(posts) >= max_posts:
                break
            last_height = new_height
        except TimeoutException:
            break
        except Exception as e:
            print(f"Error during scroll: {e}")
            break

    soup = BeautifulSoup(driver.page_source, "html.parser")

    print(f"[TAG] Found {len(soup.select(post_selector))} post elements")

    for post_elem in soup.select(post_selector):
        if len(posts) >= max_posts:
            break

        # Parse post_url + time
        a_elem = post_elem.select_one('div.isayt a.isayc[href*=".lofter.com/post/"]')
        post_url = ""
        time_text = ""
        post_time_ts = None

        if a_elem:
            post_url = a_elem["href"]
            if "title" in a_elem.attrs:
                time_match = re.search(r'(\d{2})/(\d{2})(?:\s+\d{2}:\d{2})?', a_elem["title"])
                if time_match:
                    month, day = map(int, time_match.groups())
                    year = time.localtime().tm_year  # mặc định năm hiện tại
                    try:
                        post_time_struct = time.strptime(f"{year}-{month:02d}-{day:02d}", "%Y-%m-%d")
                        post_time_ts = time.mktime(post_time_struct)
                        time_text = f"{year}-{month:02d}-{day:02d}"
                    except Exception as e:
                        print(f"[TAG] Không parse được time: {month}/{day}")

        if not post_url:
            print("[TAG] Không tìm thấy post_url hợp lệ trong post_elem, skip bài này")
            continue

        # Kiểm tra thời gian lọc
        if start_date_ts and post_time_ts and post_time_ts < start_date_ts:
            print(f"[TAG] Bỏ qua post {post_url} vì trước ngày lọc")
            stop_fetch = True
            break
        if end_date_ts and post_time_ts and post_time_ts > end_date_ts:
            print(f"[TAG] Bỏ qua post {post_url} vì sau ngày lọc")
            continue

        # Parse nội dung
        parsed = parse_by_selector(post_elem, selector_set, index_in_list=len(posts))

        # Thêm post nếu không trùng
        if post_url and post_url not in [p["url"] for p in posts]:
            posts.append({
                "title": parsed["title"],
                "url": post_url,
                "content": parsed["content"],
                "images": parsed["images"],
                "preview": parsed["content_text"][:20] + "..." if parsed["content_text"] else "",
                "time": time_text
            })
            print(f"[TAG] Thêm bài thành công: {parsed['title']} - {post_url} - {time_text}")

    # Next page nếu cần và chưa stop
    next_link = soup.select_one("a.next, a.next-page")
    if not stop_fetch and next_link and len(posts) < max_posts:
        next_url = urljoin(url, next_link["href"])
        print(f"[TAG] Chuyển sang trang tiếp theo: {next_url}")
        posts.extend(fetch_lofter_tag_posts(next_url, driver, max_posts, wait_time, continue_fetch=True, start_date_ts=start_date_ts, end_date_ts=end_date_ts, cookies_loaded=cookies_loaded))
    else:
        if stop_fetch:
            print(f"[TAG] Đã gặp bài trước ngày lọc → dừng fetch tiếp")

    print(f"[TAG] Total posts found: {len(posts)}")
    for post in posts:
        print(f"[TAG] Title: {post['title']} - URL: {post['url']} - Time: {post.get('time', '')}")

    return posts



def fetch_forum_chapters(url, driver, max_posts=50):
    current_url = url
    chapters = []
    main_title = ""

    while current_url and len(chapters) < max_posts:
        driver.get(current_url)
        soup = BeautifulSoup(driver.page_source, "html.parser")

        if not main_title:
            main_title = soup.select_one("h1.p-title-value").text.strip() if soup.select_one("h1.p-title-value") else "Forum Thread"

        print(f"[Forum] Đang fetch page: {current_url}")

        for idx, article in enumerate(soup.select("article.message")):
            if len(chapters) >= max_posts:
                break

            content_div = article.select_one("div.bbWrapper")
            user_link = article.select_one("h4.message-name a.username")

            if content_div and user_link:
                username = user_link.text.strip()
                content_html = content_div.decode_contents()
                content_text = BeautifulSoup(content_html, "html.parser").get_text(separator=" ", strip=True)
                preview = content_text[:20] + "..." if len(content_text) > 20 else content_text
                title = f"{username}: {preview}"

                chapters.append({
                    "title": title,
                    "content": content_html,
                    "images": []
                })

                print(f"[Forum] Thêm bài: {title}")

        # Xử lý next page nếu có
        next_link = soup.select_one("a.pageNav-jump--next") or soup.select_one("a.pageNavSimple-el--next")
        if next_link and len(chapters) < max_posts:
            next_href = next_link["href"]
            current_url = urljoin(current_url, next_href)
            print(f"[Forum] Chuyển sang trang tiếp theo: {current_url}")
        else:
            break

    print(f"[Forum] Total chapters found: {len(chapters)}")
    return {"main_title": main_title, "chapters": chapters}


def download_content(url, driver):
    driver.set_page_load_timeout(180)
    if not load_lofter_cookies(driver):
        print(f"Không thể tải cookie cho {url}, trả về nội dung rỗng")
        return {
            "title": f"Không tải được cookie",
            "text": "",
            "images": []
        }

    time.sleep(2)  # Để tránh load lỗi cookie xong driver.get liền

    domain = urlparse(url).netloc
    is_lofter_post = "lofter.com" in domain

    if is_lofter_post:
        selector_set = selector_configs["lofter_author"]
    else:
        selector_set = {
            "content": [
                "div.txt",
                "div.content",
                "div.text",
                "div.post-content",
                "div.g-ctc",
                "div.entry-content",
                "div.wrap + div.txt",
                "div.wrap + div.content",
                "div.wrap + div.text",
                "div.wrap",
            ],
            "images": ["img"],
            "title": ["h1", "h2", "div.ttl", "div.title"]
        }

    wait_selector = ", ".join(selector_set.get("content", ["div.g-innerbody", "div.text", "div.txtcont", "div.section", "p", "div.ct"]))

    try:
        driver.get(url)
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, wait_selector))
        )
    except TimeoutException:
        print(f"Timeout khi tải {url}, thử refresh lại lần nữa")
        try:
            driver.execute_script("window.stop();")
            time.sleep(2)
            driver.get(url)
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, wait_selector))
            )
        except TimeoutException:
            print(f"Retry cũng timeout với {url}, bỏ qua")
            return {
                "title": f"Timeout: {url}",
                "text": "",
                "images": []
            }

    time.sleep(2)
    soup = BeautifulSoup(driver.page_source, "html.parser")

    if is_lofter_post:
        post_selector = ", ".join(selector_set.get("post_list", ["div.m-post"]))
        post_elem = soup.select_one(post_selector)

        if not post_elem:
            print(f"Không tìm thấy post element với {post_selector}, thử fallback")
            for sel in selector_set.get("content", []):
                elems = soup.select(sel)
                for elem in elems:
                    # Check nếu ancestor có class m-post-about → skip
                    skip = False
                    for parent in elem.parents:
                        if parent.has_attr("class") and "m-post-about" in parent["class"]:
                            skip = True
                            break
                    if skip:
                        continue

                    # Nếu không skip → lấy elem này
                    post_elem = elem
                    break
                if post_elem:
                    break


        if not post_elem:
            print(f"Không tìm thấy post element trong {url}, thử parse body")
            post_elem = soup.select_one("body")
            if not post_elem:
                print(f"Không tìm thấy body trong {url}, skip")
                return {
                    "title": f"Không tìm thấy nội dung",
                    "text": "",
                    "images": []
                }

        parsed = parse_by_selector(post_elem, selector_set)

        if parsed["content"] == "Không tìm thấy nội dung":
            print(f"Nội dung parse không hợp lệ cho {url}, log HTML để debug")
            with open(f"temp/debug_{url.split('/')[-1]}.html", "w", encoding="utf-8") as f:
                f.write(str(soup))

    else:
        # Generic parse
        post_elem = None
        for sel in selector_set.get("content", []):
            post_elem = soup.select_one(sel)
            if post_elem:
                break

        if not post_elem:
            print(f"Không tìm thấy post element trong {url}, thử parse body")
            post_elem = soup.select_one("body")
            if not post_elem:
                print(f"Không tìm thấy body trong {url}, skip")
                return {
                    "title": f"Không tìm thấy nội dung",
                    "text": "",
                    "images": []
                }

        parsed = parse_by_selector(post_elem, selector_set)

    post_id = url.split("/")[-1]
    downloaded_images = download_images([urljoin(url, img) for img in parsed["images"]], post_id)

    content = {
        "title": parsed["title"],
        "text": parsed["content"],
        "images": downloaded_images
    }

    return content


def create_html(chapters, main_title):
    html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{main_title}</title>
    </head>
    <body>
        <h1>{main_title}</h1>
        <ul>
    """
    for idx, chapter in enumerate(chapters):
        html += f'<li><a href="#chapter{idx+1}">{chapter["title"]}</a></li>'

    html += "</ul>"

    for idx, chapter in enumerate(chapters):
        html += f'<h2 id="chapter{idx+1}">{chapter["title"]}</h2>'
        html += f'<div>{chapter["content"]}</div>'
        for img in chapter.get("images", []):
            if img and os.path.exists(img):
                img_filename = os.path.basename(img)
                shutil.copy(img, os.path.join("temp", img_filename))
                html += f'<img src="{img_filename}" alt="Image" style="max-width:100%;"><br>'

    html += "</body></html>"

    os.makedirs("temp", exist_ok=True)
    safe_main_title = sanitize_filename(main_title)
    html_file = f"temp/{safe_main_title}.html"
    with open(html_file, "w", encoding="utf-8") as f:
        f.write(html)
    return html_file

def convert_to_format(html_file, output_format, main_title):
    output_file = f"temp/{main_title}.{output_format}"

    if output_format == "epub":
        subprocess.run([
            "pandoc", html_file, "-o", output_file,
            "--metadata", f"title={main_title}",
            "--resource-path=temp/images"
        ], check=True)

    elif output_format == "pdf":
        subprocess.run([
            "wkhtmltopdf", "--enable-local-file-access", html_file, output_file
        ], check=True)

    elif output_format == "docx":
        doc = Document()
        doc.add_heading(main_title, 0)

        with open(html_file, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")

        for element in soup.find_all(["h2", "p", "img"]):
            if element.name == "h2":
                doc.add_heading(element.text, level=2)
            elif element.name == "p":
                doc.add_paragraph(element.text)
            elif element.name == "img" and element.get("src"):
                try:
                    img_path = element["src"].replace("file://", "")
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(5.0))
                except Exception as e:
                    print(f"Không thể thêm hình ảnh {element['src']}: {e}")

    elif output_format == "rar":
        output_zip = f"temp/{main_title}.zip"
        with zipfile.ZipFile(output_zip, "w") as zf:
            zf.write(html_file, arcname=os.path.basename(html_file))
            for img in os.listdir("temp/images"):
                if img.endswith(".jpg"):
                    zf.write(os.path.join("temp/images", img), arcname=img)
        os.rename(output_zip, output_file)

    else:
        raise ValueError(f"Định dạng {output_format} không được hỗ trợ.")

    return output_file

def cleanup_temp():
    if os.path.exists("temp"):
        shutil.rmtree("temp")

def sanitize_filename(filename):
    filename = unicodedata.normalize("NFD", filename)
    filename = filename.encode("ascii", "ignore").decode("utf-8")
    filename = filename.replace('đ', 'd').replace('Đ', 'D')
    filename = re.sub(r'[^A-Za-z0-9 ]', '', filename)
    filename = re.sub(r'\s+', '_', filename)
    filename = filename.strip('_')
    return filename

@app.route("/health", methods=["GET"])
def health_check():
    return jsonify({"status": "Server is running"})

@app.route("/fetch-chapters", methods=["POST"])
def fetch_chapters():
    data = request.get_json()
    url = data.get('url')
    url_type = data.get('type')
    max_posts = data.get('max_posts', 50)
    wait_time = data.get('wait_time', 1)
    continue_fetch = data.get('continue_fetch', False)
    start_date_str = data.get('start_date')
    end_date_str = data.get('end_date')

    start_date_ts = None
    end_date_ts = None

    if start_date_str:
        start_date_ts = time.mktime(time.strptime(start_date_str, "%Y-%m-%d"))
    if end_date_str:
        end_date_ts = time.mktime(time.strptime(end_date_str, "%Y-%m-%d"))
    driver = setup_selenium()

    try:
        if url_type == "ao3_tag":
            items = fetch_ao3_works(url, driver)
            return jsonify(items)
        elif url_type == "lofter_author":
            items = fetch_lofter_posts(url, driver, False, max_posts, wait_time, continue_fetch, start_date_ts, end_date_ts)
            if not items:
                return jsonify({"error": "Không tìm thấy bài viết"}), 400
            return jsonify(items)
        elif url_type == "lofter_tag":
            items = fetch_lofter_tag_posts(url, driver, max_posts, wait_time, continue_fetch, start_date_ts, end_date_ts)
            if not items:
                return jsonify({"error": "Không tìm thấy bài viết"}), 400
            return jsonify(items)
        elif url_type == "forum":
            result = fetch_forum_chapters(url, driver)
            return jsonify(result)
        else:
            return jsonify({"error": "Invalid URL type"}), 400
    finally:
        pass

@app.route("/download", methods=["POST"])
def download():
    data = request.get_json()
    driver = setup_selenium()

    try:
        type_ = data.get("type", "generic")
        output_format = data.get("format", "epub").lower()
        main_title = data.get("main_title", "Downloaded Content")
        safe_main_title = sanitize_filename(main_title)

        if type_ == "forum":
            chapters = data.get("chapters", [])
            html_file = create_html(chapters, safe_main_title)
            output_file = convert_to_format(html_file, output_format, safe_main_title)

        elif type_ == "lofter":
            urls = data.get("urls", [])
            chapters = []
            
            driver = setup_selenium()
            load_lofter_cookies(driver)

            for url in urls:
                print(f"Download bằng driver global: {url}")
                time.sleep(1.5)

                try:
                    # Không cần load_lofter_cookies(driver) nữa nếu đã load bên ngoài rồi.
                    content = download_content(url, driver)  # <<< xài driver global
                    chapters.append({
                        "title": content["title"],
                        "content": content["text"],
                        "images": content["images"]
                    })
                except Exception as e:
                    print(f"Error downloading {url}: {e}")
                    chapters.append({
                        "title": f"Error downloading {url}",
                        "content": "",
                        "images": []
                    })
                finally:
                    time.sleep(1.5)

            html_file = create_html(chapters, safe_main_title)
            output_file = convert_to_format(html_file, output_format, safe_main_title)

        elif type_ == "ao3":
            urls = data.get("urls", [])
            chapters = []
            for url in urls:
                driver.get(url)
                time.sleep(2)
                soup = BeautifulSoup(driver.page_source, "html.parser")
                content_elem = soup.select_one("div#workskin")
                content_html = content_elem.decode_contents() if content_elem else ""
                content_text = content_elem.get_text(separator=" ", strip=True) if content_elem else ""
                chapters.append({
                    "title": soup.select_one("h2.title").text.strip() if soup.select_one("h2.title") else url,
                    "content": content_html,
                    "images": []
                })
            html_file = create_html(chapters, safe_main_title)
            output_file = convert_to_format(html_file, output_format, safe_main_title)

        else:
            urls = data.get("urls", [])
            chapters = []
            for url in urls:
                content = download_content(url, driver)
                chapters.append({
                    "title": content["title"],
                    "content": content["text"],
                    "images": content["images"]
                })
            html_file = create_html(chapters, safe_main_title)
            output_file = convert_to_format(html_file, output_format, safe_main_title)

        with open(output_file, "rb") as f:
            file_data = f.read()
        file_name = os.path.basename(output_file)
        cleanup_temp()
        return send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            download_name=file_name,
            mimetype=f"application/{'vnd.openxmlformats-officedocument.wordprocessingml.document' if output_format == 'docx' else output_format}"
        )
    finally:
        pass

@app.route('/')
def serve_index():
    return send_from_directory('.', 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('.', path)

if __name__ == "__main__":
    try:
        app.run(debug=True, host="0.0.0.0", port=5000, threaded=True)
    finally:
        if driver:
            driver.quit()